"""
Convert a Substack post to a .docx file.

Usage:
    python convert.py <url> [--out OUTPUT.docx]

Preserves inline bold/italic/links, headings, blockquotes, lists, and images
(downloaded and embedded). Body text is Times New Roman 12pt.
"""

from __future__ import annotations

import argparse
import io
import re
import sys
from pathlib import Path
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0 Safari/537.36"
    )
}

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)
MAX_IMAGE_WIDTH_IN = 6.0  # inches — fits inside 1" margins on letter paper


def fetch(url: str) -> str:
    r = requests.get(url, headers=HEADERS, timeout=30)
    r.raise_for_status()
    return r.text


def find_article(soup: BeautifulSoup) -> Tag:
    # Substack renders the body inside .available-content (with .body.markup inside).
    for selector in [
        "div.available-content",
        "div.body.markup",
        "article",
    ]:
        node = soup.select_one(selector)
        if node:
            return node
    raise RuntimeError("Could not locate article body in page.")


def extract_title(soup: BeautifulSoup) -> str:
    for selector in ["h1.post-title", "h1", 'meta[property="og:title"]']:
        node = soup.select_one(selector)
        if node:
            if node.name == "meta":
                return node.get("content", "").strip()
            return node.get_text(strip=True)
    return "Untitled"


def extract_subtitle(soup: BeautifulSoup) -> str | None:
    for selector in ["h3.subtitle", "h2.subtitle", 'meta[property="og:description"]']:
        node = soup.select_one(selector)
        if node:
            text = node.get("content", "").strip() if node.name == "meta" else node.get_text(strip=True)
            if text:
                return text
    return None


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE


def add_runs(paragraph, node, *, bold=False, italic=False, link=False) -> None:
    """Walk inline children and append runs with cumulative formatting."""
    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        run = paragraph.add_run(text)
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if link:
            run.underline = True
        return

    if not isinstance(node, Tag):
        return

    tag = node.name.lower()
    b = bold or tag in ("strong", "b")
    i = italic or tag in ("em", "i")
    lk = link or tag == "a"

    if tag == "br":
        paragraph.add_run().add_break()
        return

    for child in node.children:
        add_runs(paragraph, child, bold=b, italic=i, link=lk)


def download_image(src: str) -> bytes | None:
    try:
        r = requests.get(src, headers=HEADERS, timeout=30)
        r.raise_for_status()
        return r.content
    except Exception as e:
        print(f"  ! image failed ({src}): {e}", file=sys.stderr)
        return None


def add_image(document: Document, src: str) -> None:
    data = download_image(src)
    if not data:
        return
    try:
        img = Image.open(io.BytesIO(data))
        img.load()
    except Exception as e:
        print(f"  ! could not decode image ({src}): {e}", file=sys.stderr)
        return

    # python-docx wants a file-like object; re-encode to a normalized PNG/JPEG
    # so animated or exotic formats don't blow up.
    buf = io.BytesIO()
    fmt = "PNG" if img.mode in ("RGBA", "LA", "P") else "JPEG"
    if fmt == "JPEG" and img.mode != "RGB":
        img = img.convert("RGB")
    img.save(buf, format=fmt)
    buf.seek(0)

    px_w = img.width
    # 96 dpi assumption for web images; cap at MAX_IMAGE_WIDTH_IN.
    width_in = min(MAX_IMAGE_WIDTH_IN, px_w / 96.0)
    document.add_picture(buf, width=Inches(width_in))


def process_block(document: Document, node: Tag) -> None:
    tag = node.name.lower() if node.name else ""

    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(tag[1])
        p = document.add_heading(level=min(level, 4))
        for child in node.children:
            add_runs(p, child, bold=True)
        for run in p.runs:
            run.font.name = BODY_FONT
        return

    if tag == "p":
        p = document.add_paragraph()
        for child in node.children:
            add_runs(p, child)
        return

    if tag == "blockquote":
        for child_block in node.find_all(["p"], recursive=False) or [node]:
            p = document.add_paragraph(style="Intense Quote" if "Intense Quote" in [s.name for s in document.styles] else None)
            for child in child_block.children:
                add_runs(p, child, italic=True)
        return

    if tag in ("ul", "ol"):
        style = "List Bullet" if tag == "ul" else "List Number"
        for li in node.find_all("li", recursive=False):
            p = document.add_paragraph(style=style)
            for child in li.children:
                add_runs(p, child)
        return

    if tag == "hr":
        document.add_paragraph("―" * 20)
        return

    if tag == "figure":
        img = node.find("img")
        if img and img.get("src"):
            add_image(document, img["src"])
        cap = node.find("figcaption")
        if cap:
            p = document.add_paragraph()
            for child in cap.children:
                add_runs(p, child, italic=True)
        return

    if tag == "img" and node.get("src"):
        add_image(document, node["src"])
        return

    # Substack image containers: only trigger for nodes whose own class
    # signals an image wrapper — never for generic divs that merely
    # contain an image somewhere deep inside.
    classes = " ".join(node.get("class", [])) if isinstance(node, Tag) else ""
    if any(k in classes for k in ("captioned-image", "image-link", "image2")):
        img = node.find("img")
        if img and img.get("src"):
            add_image(document, img["src"])
            cap = node.find("figcaption")
            if cap:
                p = document.add_paragraph()
                for child in cap.children:
                    add_runs(p, child, italic=True)
            return

    # Fallback: recurse into children looking for known blocks.
    for child in node.children:
        if isinstance(child, Tag):
            process_block(document, child)


def slugify(text: str) -> str:
    text = re.sub(r"[^\w\s-]", "", text).strip().lower()
    return re.sub(r"[-\s]+", "-", text)[:80] or "post"


def convert(url: str, out_path: Path) -> Path:
    print(f"Fetching {url}")
    html = fetch(url)
    soup = BeautifulSoup(html, "html.parser")

    title = extract_title(soup)
    subtitle = extract_subtitle(soup)
    article = find_article(soup)

    document = Document()
    set_default_font(document)

    # Title
    t = document.add_paragraph()
    run = t.add_run(title)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(18)

    if subtitle:
        s = document.add_paragraph()
        run = s.add_run(subtitle)
        run.italic = True
        run.font.name = BODY_FONT
        run.font.size = Pt(13)

    document.add_paragraph()  # spacer

    # Walk top-level article children as blocks.
    top_children = [c for c in article.children if isinstance(c, Tag)]
    if not top_children:
        # Some Substack pages nest one level deeper.
        inner = article.find("div", class_="body")
        if inner:
            top_children = [c for c in inner.children if isinstance(c, Tag)]

    for node in top_children:
        process_block(document, node)

    if out_path.is_dir() or out_path.suffix == "":
        out_path = out_path / f"{slugify(title)}.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)
    print(f"Wrote {out_path}")
    return out_path


def main() -> None:
    ap = argparse.ArgumentParser(description="Convert a Substack post to .docx")
    ap.add_argument("url", help="Substack post URL")
    ap.add_argument(
        "--out",
        default=str(Path.cwd()),
        help="Output file or directory (default: current working directory)",
    )
    args = ap.parse_args()

    out = Path(args.out).expanduser().resolve()
    convert(args.url, out)


if __name__ == "__main__":
    main()
